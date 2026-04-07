from __future__ import annotations

import hashlib
import io
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import List

import av
import streamlit as st
import whisper
from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel
from streamlit_webrtc import WebRtcMode, webrtc_streamer


@dataclass
class Chunk:
    source_name: str
    page_number: int
    heading: str
    text: str


@dataclass
class LocalDocument:
    name: str
    content: bytes


STOP_WORDS = {
    "the", "a", "an", "for", "to", "of", "in", "on", "and", "or", "is", "are",
    "what", "how", "when", "where", "show", "give", "need", "about", "before",
    "after", "with", "from", "into", "procedure", "procedures", "step", "steps",
    "manual", "manuals", "please", "find", "tell", "me",
}

PROCEDURE_TERMS = {
    "procedure", "procedures", "step", "steps", "test", "testing", "check",
    "checks", "inspection", "inspect", "method", "operation", "operational",
    "startup", "shutdown", "pre-operational", "preoperational", "checklist",
    "instructions", "verification",
}

GENERIC_TERMS = {
    "description", "overview", "general", "introduction", "scope", "purpose",
    "background", "summary", "definitions",
}


def normalize_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def slug_terms(value: str) -> List[str]:
    return [term for term in re.findall(r"[a-z0-9][a-z0-9\-]+", value.lower()) if term]


def is_heading(line: str) -> bool:
    cleaned = line.strip()
    if len(cleaned) < 4 or len(cleaned) > 120:
        return False

    heading_patterns = [
        r"^\d+(\.\d+)*[\)\.]?\s+[A-Z][A-Za-z0-9 ,:/\-]{2,}$",
        r"^[A-Z][A-Z0-9 ,:/\-\(\)]{4,}$",
        r"^(Procedure|Procedures|Steps|Instructions|Method|Operation|Inspection)\b",
    ]
    return any(re.match(pattern, cleaned) for pattern in heading_patterns)


def split_page_into_chunks(source_name: str, page_number: int, page_text: str) -> List[Chunk]:
    text = normalize_text(page_text)
    if not text:
        return []

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections: List[Chunk] = []
    current_heading = f"Page {page_number}"
    buffer: List[str] = []

    def flush_buffer() -> None:
        nonlocal buffer
        if not buffer:
            return
        section_text = normalize_text("\n".join(buffer))
        if section_text:
            sections.append(
                Chunk(
                    source_name=source_name,
                    page_number=page_number,
                    heading=current_heading,
                    text=section_text,
                )
            )
        buffer = []

    for line in lines:
        if is_heading(line):
            flush_buffer()
            current_heading = line
            continue
        buffer.append(line)

    flush_buffer()

    expanded_sections: List[Chunk] = []
    for section in sections:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", section.text) if p.strip()]
        if not paragraphs:
            continue

        current_parts: List[str] = []
        current_length = 0
        for paragraph in paragraphs:
            projected = current_length + len(paragraph)
            if current_parts and projected > 1200:
                expanded_sections.append(
                    Chunk(
                        source_name=section.source_name,
                        page_number=section.page_number,
                        heading=section.heading,
                        text="\n\n".join(current_parts),
                    )
                )
                overlap = current_parts[-1:]
                current_parts = overlap + [paragraph]
                current_length = sum(len(item) for item in current_parts)
            else:
                current_parts.append(paragraph)
                current_length = projected

        if current_parts:
            expanded_sections.append(
                Chunk(
                    source_name=section.source_name,
                    page_number=section.page_number,
                    heading=section.heading,
                    text="\n\n".join(current_parts),
                )
            )

    return expanded_sections


def split_document_text_into_chunks(source_name: str, text: str) -> List[Chunk]:
    return split_page_into_chunks(
        source_name=source_name,
        page_number=1,
        page_text=text,
    )


def load_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def extract_chunks_from_file(name: str, file_bytes: bytes) -> List[Chunk]:
    file_name = name.lower()
    chunks: List[Chunk] = []

    if file_name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            chunks.extend(
                split_page_into_chunks(
                    source_name=name,
                    page_number=page_index,
                    page_text=page_text,
                )
            )
    elif file_name.endswith(".docx"):
        docx_text = load_docx_text(file_bytes)
        chunks.extend(
            split_document_text_into_chunks(
                source_name=name,
                text=docx_text,
            )
        )

    return chunks


def load_document_chunks(uploaded_files) -> List[Chunk]:
    all_chunks: List[Chunk] = []
    for uploaded_file in uploaded_files:
        if hasattr(uploaded_file, "getvalue"):
            content = uploaded_file.getvalue()
        else:
            content = uploaded_file.content
        all_chunks.extend(extract_chunks_from_file(uploaded_file.name, content))

    return [chunk for chunk in all_chunks if chunk.text.strip()]


def load_local_default_documents() -> List[LocalDocument]:
    default_dir = Path(__file__).with_name("default_manuals")
    if not default_dir.exists():
        return []

    documents: List[LocalDocument] = []
    for path in sorted(default_dir.iterdir()):
        if path.is_file() and path.suffix.lower() in {".pdf", ".docx"}:
            documents.append(LocalDocument(name=path.name, content=path.read_bytes()))
    return documents


def fingerprint_files(uploaded_files) -> str:
    digest = hashlib.sha256()
    for uploaded_file in uploaded_files:
        name = uploaded_file.name
        if hasattr(uploaded_file, "getvalue"):
            content = uploaded_file.getvalue()
            size = getattr(uploaded_file, "size", len(content))
        else:
            content = uploaded_file.content
            size = len(content)

        digest.update(name.encode("utf-8"))
        digest.update(str(size).encode("utf-8"))
        digest.update(content)
    return digest.hexdigest()


@st.cache_resource(show_spinner=False)
def load_whisper_model(model_name: str):
    return whisper.load_model(model_name)


def transcribe_audio_bytes(audio_bytes: bytes, model_name: str) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
        temp_audio.write(audio_bytes)
        temp_path = temp_audio.name

    try:
        model = load_whisper_model(model_name)
        result = model.transcribe(temp_path, language="en", fp16=False)
        return (result.get("text") or "").strip()
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def build_search_index(chunks: List[Chunk]):
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    matrix = vectorizer.fit_transform([f"{chunk.heading}\n{chunk.text}" for chunk in chunks])
    return vectorizer, matrix


def search_chunks(question: str, vectorizer, matrix, chunks: List[Chunk], top_k: int = 3):
    if matrix.shape[0] == 0:
        return []

    query_vector = vectorizer.transform([question])
    scores = linear_kernel(query_vector, matrix).flatten()
    ranked_indices = scores.argsort()[::-1]

    quoted_phrases = [match.strip() for match in re.findall(r'"([^"]+)"', question.lower()) if match.strip()]
    capitalized_phrases = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b", question)
    detected_phrases = quoted_phrases + [phrase.lower() for phrase in capitalized_phrases]
    query_terms = [term for term in slug_terms(question) if term not in STOP_WORDS and len(term) > 2]
    important_terms = [term for term in query_terms if term not in PROCEDURE_TERMS]

    rescored = []
    for idx in ranked_indices[: min(len(chunks), 20)]:
        base_score = float(scores[idx])
        if base_score <= 0:
            continue

        chunk = chunks[idx]
        haystack = f"{chunk.heading}\n{chunk.text}".lower()
        heading = chunk.heading.lower()
        score = base_score

        procedure_boost = sum(1 for term in PROCEDURE_TERMS if term in heading) * 0.12
        score += procedure_boost

        generic_penalty = sum(1 for term in GENERIC_TERMS if term in heading) * 0.06
        score -= generic_penalty

        phrase_hits = 0
        for phrase in detected_phrases:
            if len(phrase.split()) >= 2 and phrase in haystack:
                phrase_hits += 1
        score += phrase_hits * 0.35

        term_hits = sum(1 for term in important_terms if term in haystack)
        if important_terms:
            coverage = term_hits / max(len(important_terms), 1)
            score += coverage * 0.45

        if any(term in heading for term in {"test", "procedure", "steps", "inspection", "check"}):
            score += 0.15

        rescored.append((chunk, score))

    rescored.sort(key=lambda item: item[1], reverse=True)
    results = []
    for chunk, score in rescored[:top_k]:
        if score <= 0:
            continue
        results.append((chunk, score))
    return results


def ensure_session_defaults() -> None:
    defaults = {
        "manual_fingerprint": None,
        "chunks": [],
        "vectorizer": None,
        "matrix": None,
        "transcript": "",
        "live_transcript": "",
        "live_audio_buffer": bytearray(),
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def append_audio_frames_to_buffer(frames: List[av.AudioFrame]) -> None:
    if not frames:
        return

    with io.BytesIO() as wav_buffer:
        import wave

        first = frames[0].to_ndarray()
        sample_rate = frames[0].sample_rate
        channels = 1 if first.ndim == 1 else first.shape[0]

        with wave.open(wav_buffer, "wb") as wav_file:
            wav_file.setnchannels(channels)
            wav_file.setsampwidth(2)
            wav_file.setframerate(sample_rate)

            for frame in frames:
                array = frame.to_ndarray()
                if array.ndim > 1:
                    array = array.T
                wav_file.writeframes(array.astype("int16").tobytes())

        st.session_state.live_audio_buffer.extend(wav_buffer.getvalue())


def process_live_audio_if_available(ctx, model_name: str) -> None:
    if not ctx or not ctx.state.playing or not ctx.audio_receiver:
        return

    try:
        frames = ctx.audio_receiver.get_frames(timeout=0.2)
    except Exception:
        frames = []

    if not frames:
        return

    append_audio_frames_to_buffer(frames)

    if len(st.session_state.live_audio_buffer) > 240000:
        transcript = transcribe_audio_bytes(bytes(st.session_state.live_audio_buffer), model_name)
        if transcript:
            st.session_state.live_transcript = transcript


def index_manuals(uploaded_files) -> None:
    chunks = load_document_chunks(uploaded_files)
    if not chunks:
        raise ValueError("Yuklenen dosyalardan okunabilir metin cikarilamadi.")

    vectorizer, matrix = build_search_index(chunks)
    st.session_state.manual_fingerprint = fingerprint_files(uploaded_files)
    st.session_state.chunks = chunks
    st.session_state.vectorizer = vectorizer
    st.session_state.matrix = matrix


def main() -> None:
    st.set_page_config(page_title="Manual Audit Assistant", page_icon="📘", layout="wide")
    ensure_session_defaults()

    st.markdown(
        """
        <style>
        .stApp {
            background:
                radial-gradient(circle at top right, rgba(14, 116, 144, 0.22), transparent 28%),
                radial-gradient(circle at left top, rgba(30, 64, 175, 0.16), transparent 24%),
                linear-gradient(180deg, #08111f 0%, #0e1726 100%);
        }
        .hero-box {
            padding: 28px 32px;
            border-radius: 22px;
            background: linear-gradient(135deg, rgba(15, 23, 42, 0.94), rgba(12, 74, 110, 0.90));
            border: 1px solid rgba(125, 211, 252, 0.18);
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.28);
            margin-bottom: 18px;
        }
        .hero-kicker {
            display: inline-block;
            padding: 6px 12px;
            border-radius: 999px;
            background: rgba(125, 211, 252, 0.12);
            color: #bae6fd;
            font-size: 0.85rem;
            letter-spacing: 0.06em;
            text-transform: uppercase;
            margin-bottom: 12px;
        }
        .hero-title {
            font-size: 3rem;
            font-weight: 800;
            line-height: 1.05;
            color: #f8fafc;
            margin: 0 0 12px 0;
        }
        .hero-text {
            font-size: 1.05rem;
            color: #dbeafe;
            max-width: 900px;
            margin-bottom: 0;
        }
        .dedication-box {
            margin-top: 18px;
            padding: 18px 20px;
            border-radius: 18px;
            background: linear-gradient(135deg, rgba(30, 41, 59, 0.92), rgba(8, 47, 73, 0.90));
            border: 1px solid rgba(186, 230, 253, 0.18);
        }
        .dedication-title {
            font-size: 1.05rem;
            color: #93c5fd;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            margin-bottom: 8px;
        }
        .dedication-name {
            font-size: 2.2rem;
            font-weight: 800;
            color: #f8fafc;
            margin: 0;
        }
        .info-strip {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 12px;
            margin: 18px 0 10px 0;
        }
        .info-card {
            background: rgba(15, 23, 42, 0.72);
            border: 1px solid rgba(125, 211, 252, 0.12);
            border-radius: 16px;
            padding: 14px 16px;
        }
        .info-label {
            color: #93c5fd;
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            margin-bottom: 6px;
        }
        .info-value {
            color: #f8fafc;
            font-size: 1rem;
            font-weight: 700;
        }
        .section-card {
            background: rgba(15, 23, 42, 0.72);
            border: 1px solid rgba(148, 163, 184, 0.16);
            border-radius: 20px;
            padding: 18px 18px 10px 18px;
            margin-bottom: 14px;
            box-shadow: 0 14px 34px rgba(0, 0, 0, 0.18);
        }
        @media (max-width: 900px) {
            .hero-title { font-size: 2.1rem; }
            .dedication-name { font-size: 1.6rem; }
            .info-strip { grid-template-columns: 1fr; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="hero-box">
            <div class="hero-kicker">Audit Intelligence Workspace</div>
            <div class="hero-title">Manual Audit Assistant</div>
            <p class="hero-text">
                PDF ve Word manuellerinizi yukleyin, Ingilizce denetci sorusunu mikrofondan alin
                ve sadece ilgili prosedur bolumunu hizli, temiz ve kaynak bilgisiyle birlikte bulun.
            </p>
            <div class="dedication-box">
                <div class="dedication-title">Ithafen</div>
                <p class="dedication-name">Kaptan Onur Sonmez</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="info-strip">
            <div class="info-card">
                <div class="info-label">Desteklenen Dosyalar</div>
                <div class="info-value">PDF ve Word (.docx)</div>
            </div>
            <div class="info-card">
                <div class="info-label">Ses Girdisi</div>
                <div class="info-value">Whisper ile Ingilizce transkript</div>
            </div>
            <div class="info-card">
                <div class="info-label">Cikti</div>
                <div class="info-value">Sadece ilgili prosedur bolumu</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    default_documents = load_local_default_documents()

    with st.sidebar:
        st.header("Ayarlar")
        whisper_model_name = st.selectbox(
            "Whisper modeli",
            options=["base.en", "small.en", "medium.en"],
            index=0,
            help="base.en CPU icin daha hafif, medium.en daha dogru ama daha yavas olabilir.",
        )
        st.caption("Whisper icin sisteminizde ffmpeg kurulu olmali.")
        listening_mode = st.radio(
            "Dinleme modu",
            options=["Kayit modu", "Canli dinleme"],
            index=0,
            help="Canli dinleme tarayicidan mikrofon akisini baslatir.",
        )
        use_default_docs = st.toggle(
            "Varsayilan manuelleri kullan",
            value=bool(default_documents),
            disabled=not bool(default_documents),
            help="Repo icindeki default_manuals klasorundeki PDF ve Word dosyalarini otomatik kullanir.",
        )
        if default_documents:
            st.caption(f"Hazir dokuman sayisi: {len(default_documents)}")
        else:
            st.caption("Varsayilan manuel yok. default_manuals klasorune PDF veya DOCX ekleyebilirsiniz.")
        st.caption("Arama motoru prosedur ve test odakli yeniden siralama ile guclendirildi.")

    uploaded_files = st.file_uploader(
        "PDF veya Word manuellerini yukleyin",
        type=["pdf", "docx"],
        accept_multiple_files=True,
    )

    selected_documents = list(uploaded_files) if uploaded_files else []
    if use_default_docs and default_documents:
        selected_documents.extend(default_documents)

    if selected_documents:
        current_fingerprint = fingerprint_files(selected_documents)
        if st.session_state.manual_fingerprint != current_fingerprint:
            try:
                with st.spinner("Dosyalar indeksleniyor..."):
                    index_manuals(selected_documents)
                st.success(f"{len(selected_documents)} dosya islendi, arama icin hazir.")
            except Exception as exc:
                st.session_state.manual_fingerprint = None
                st.session_state.chunks = []
                st.session_state.vectorizer = None
                st.session_state.matrix = None
                st.error(f"Dosyalar islenemedi: {exc}")
    else:
        st.info("Devam etmek icin en az bir PDF veya Word dosyasi yukleyin.")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader("1. Mikrofonu dinle")
        if listening_mode == "Kayit modu":
            audio_file = st.audio_input(
                "Ingilizce soruyu kaydedin",
                sample_rate=16000,
                help="Tarayici mikrofon izni isteyecektir. Kayit WAV olarak alinir.",
            )

            if audio_file is not None:
                st.audio(audio_file)
                if st.button("Konusmayi metne cevir", type="primary", use_container_width=True):
                    with st.spinner("Whisper ile metne cevriliyor..."):
                        transcript = transcribe_audio_bytes(audio_file.getvalue(), whisper_model_name)
                    st.session_state.transcript = transcript
        else:
            live_ctx = webrtc_streamer(
                key="live-audio",
                mode=WebRtcMode.SENDONLY,
                media_stream_constraints={"video": False, "audio": True},
                rtc_configuration={"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]},
                audio_receiver_size=256,
            )
            process_live_audio_if_available(live_ctx, whisper_model_name)
            st.caption("START ile mikrofon akisina baslayin. Akis geldikce transkript guncellenir.")
            if st.button("Canli transkripti temizle", use_container_width=True):
                st.session_state.live_audio_buffer = bytearray()
                st.session_state.live_transcript = ""
            if st.session_state.live_transcript:
                st.markdown("### Canli transkript")
                st.write(st.session_state.live_transcript)

        current_question = st.session_state.live_transcript or st.session_state.transcript
        question_text = st.text_area(
            "2. Denetci sorusu",
            value=current_question,
            height=180,
            placeholder="Orn: What is the inspection procedure before restarting the pump?",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader("3. Ilgili proseduru bul")
        can_search = all(
            [
                st.session_state.vectorizer is not None,
                st.session_state.matrix is not None,
                bool(question_text.strip()),
            ]
        )

        if st.button("Manuallerde ara", disabled=not can_search, use_container_width=True):
            results = search_chunks(
                question=question_text,
                vectorizer=st.session_state.vectorizer,
                matrix=st.session_state.matrix,
                chunks=st.session_state.chunks,
                top_k=3,
            )

            if not results:
                st.warning("Ilgili bir prosedur bolumu bulunamadi. Soruyu daha spesifik yazmayi deneyin.")
            else:
                best_chunk, best_score = results[0]
                st.success("En ilgili prosedur bolumu bulundu.")
                st.markdown("### En ilgili prosedur")
                st.caption(
                    f"Kaynak: {best_chunk.source_name} | Sayfa: {best_chunk.page_number} | Bolum: {best_chunk.heading} | Benzerlik: {best_score:.3f}"
                )
                st.markdown(best_chunk.text)

                if len(results) > 1:
                    with st.expander("Yakin diger eslesmeler"):
                        for chunk, score in results[1:]:
                            st.markdown(
                                f"**{chunk.source_name}** | Sayfa **{chunk.page_number}** | Bolum **{chunk.heading}** | Skor **{score:.3f}**"
                            )
                            st.write(chunk.text)
                            st.divider()

        if st.session_state.transcript:
            st.markdown("### Son transkript")
            st.write(st.session_state.transcript)
        st.markdown("</div>", unsafe_allow_html=True)

    st.divider()
    st.caption(
        "Not: Bu surum PDF ve Word metnini cikarip TF-IDF ile en ilgili prosedur bolumunu getirir. Gorsel tabanli PDF'lerde OCR gerekebilir."
    )


if __name__ == "__main__":
    main()
